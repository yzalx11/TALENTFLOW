# app/utils/file_parser.py
"""
文档解析工厂 — 统一入口 parse_resume_file() 按扩展名路由到独立解析器，
支持 PDF / DOCX / TXT / DOC / 图片 OCR，PDF 扫描件自动回退 OCR。
"""
import os
import shutil
import subprocess
import time
import pdfplumber
from docx import Document
from fastapi import UploadFile

from app.core.config import settings
from app.core.logger import logger

# OCR 依赖（可选，未安装时降级提示）
try:
    import pytesseract
    from PIL import Image
    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    _PDF2IMAGE_AVAILABLE = True
except ImportError:
    _PDF2IMAGE_AVAILABLE = False

# ============================================================
# 配置常量
# ============================================================
UPLOAD_DIR = settings.UPLOAD_DIR
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 支持的文件格式
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt",
                        ".png", ".jpg", ".jpeg", ".bmp", ".tiff"}


# ============================================================
# 文件落盘
# ============================================================

def save_upload_file_to_disk(file: UploadFile) -> str:
    """将上传的文件保存到本地物理硬盘，返回文件路径"""
    safe_filename = f"{int(time.time())}_{file.filename}"
    file_location = os.path.join(UPLOAD_DIR, safe_filename)

    try:
        with open(file_location, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        return file_location
    except Exception as e:
        logger.error(f"❌ 文件落盘失败: {e}")
        return ""
    finally:
        file.file.seek(0)


# ============================================================
# 工厂入口
# ============================================================

# 显式路由表 — 扩展名 → 解析函数，静态分析可追踪
_DISPATCH = {}  # 在各解析函数定义后填充，见文件末尾


def parse_resume_file(file_location: str, fallback_text: str = "") -> str:
    """
    统一简历解析入口 — 按文件扩展名分发到对应的独立解析器。

    支持格式: PDF, DOCX, TXT
    预留格式: DOC (旧版 Word), PNG/JPG/BMP/TIFF (图片 OCR)
    不支持的格式返回 fallback_text
    """
    if not file_location or not os.path.exists(file_location):
        return fallback_text

    file_ext = os.path.splitext(file_location)[1].lower()

    parse_func = _DISPATCH.get(file_ext)
    if parse_func is None:
        # 检查是否为已声明的保留格式（解析器尚未实现）
        if file_ext in SUPPORTED_EXTENSIONS:
            logger.warning(f"⚠️ 格式 {file_ext} 已声明但解析器未注册")
        return fallback_text

    try:
        return parse_func(file_location)
    except NotImplementedError as e:
        logger.warning(f"⚠️ {e}")
        return fallback_text
    except Exception as e:
        logger.error(f"❌ 文件 {file_location} 解析失败: {e}")
        return fallback_text


# ============================================================
# 各格式独立解析器
# ============================================================

def _parse_pdf(file_location: str) -> str:
    """PDF 解析：先文本层提取，扫描件自动回退 OCR"""
    return _parse_pdf_with_ocr_fallback(file_location)


def _parse_docx(file_location: str) -> str:
    """DOCX 解析：遍历段落 + 表格"""
    doc = Document(file_location)
    full_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text_parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                [cell.text.strip() for cell in row.cells if cell.text.strip()]
            )
            if row_text:
                full_text_parts.append(row_text)

    result = "\n".join(full_text_parts).strip()
    if not result:
        raise ValueError("DOCX 未提取到任何文本内容")
    return result


def _parse_txt(file_location: str) -> str:
    """TXT 解析：直接读取"""
    with open(file_location, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()

    if not text:
        raise ValueError("TXT 文件内容为空")
    return text


def _parse_doc(file_location: str) -> str:
    """
    旧版 .doc 格式解析 — 通过外部工具转换后提取文本。

    尝试顺序:
      1. antiword (轻量，Linux/Mac 友好)
      2. LibreOffice 无头模式 (跨平台，功能完整)

    如果两者都不可用，抛出 NotImplementedError 并提示安装方案。
    """
    # 方案1: antiword
    try:
        result = subprocess.run(
            ["antiword", file_location],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except FileNotFoundError:
        pass
    except Exception as e:
        logger.warning(f"antiword 转换失败: {e}")

    # 方案2: LibreOffice 无头模式
    try:
        output_dir = os.path.dirname(file_location)
        result = subprocess.run(
            [
                "libreoffice", "--headless", "--convert-to", "txt",
                "--outdir", output_dir, file_location
            ],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            base_name = os.path.splitext(os.path.basename(file_location))[0]
            txt_path = os.path.join(output_dir, f"{base_name}.txt")
            if os.path.exists(txt_path):
                with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read().strip()
                os.remove(txt_path)  # 清理临时文件
                if text:
                    return text
    except FileNotFoundError:
        pass
    except Exception as e:
        logger.warning(f"LibreOffice 转换失败: {e}")

    raise NotImplementedError(
        "旧版 .doc 格式需要安装 antiword 或 LibreOffice 才能解析。\n"
        "  - Ubuntu/Debian: sudo apt install antiword\n"
        "  - macOS: brew install antiword\n"
        "  - 通用方案: 安装 LibreOffice 后重试"
    )


def _ocr_image(file_path: str) -> str:
    """单张图片 OCR，返回提取文本"""
    if not _OCR_AVAILABLE:
        raise NotImplementedError(
            "OCR 未启用：需安装 pytesseract + Pillow + 系统 tesseract-ocr"
        )
    img = Image.open(file_path)
    # 中文 + 英文混合识别
    text = pytesseract.image_to_string(img, lang="chi_sim+eng")
    return text.strip()


def _parse_image(file_location: str) -> str:
    """图片简历 OCR 解析 — pytesseract 实现"""
    text = _ocr_image(file_location)
    if not text:
        raise ValueError("OCR 未提取到任何文本内容（图片可能为空白或清晰度过低）")
    return text


def _parse_pdf_with_ocr_fallback(file_location: str) -> str:
    """
    PDF 解析增强版：先走 pdfplumber 提取文本层；
    若未提取到文字（扫描件），则逐页转图片后 OCR。
    """
    # 第一轮：普通 PDF 文本提取
    with pdfplumber.open(file_location) as pdf:
        full_text_parts = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                full_text_parts.append(text)
            tables = page.extract_tables()
            for table in tables:
                if table:
                    for row in table:
                        row_text = " | ".join(
                            [cell if cell is not None else "" for cell in row]
                        )
                        if row_text.strip():
                            full_text_parts.append(row_text)

    result = "\n".join(full_text_parts).strip()
    if result:
        return result  # 普通 PDF，直接返回

    # 第二轮：OCR 兜底 (扫描件 PDF)
    if not _PDF2IMAGE_AVAILABLE or not _OCR_AVAILABLE:
        raise NotImplementedError(
            "PDF 扫描件需要安装 pdf2image + poppler-utils + tesseract-ocr 才能 OCR 识别。\n"
            "  - Ubuntu/Debian: sudo apt install poppler-utils tesseract-ocr tesseract-ocr-chi-sim\n"
            "  - macOS: brew install poppler tesseract tesseract-lang"
        )

    logger.info(f"📷 PDF 文本层为空，启用 OCR 扫描件模式: {os.path.basename(file_location)}")
    ocr_parts = []
    images = convert_from_path(file_location, dpi=300)
    for i, img in enumerate(images):
        page_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        if page_text.strip():
            ocr_parts.append(page_text.strip())
        logger.info(f"  OCR 第 {i+1}/{len(images)} 页 → {len(page_text.strip())} 字符")

    result = "\n".join(ocr_parts).strip()
    if not result:
        raise ValueError("PDF 扫描件 OCR 未提取到任何文本内容")
    return result


# ============================================================
# 路由表注册 — 各解析函数定义完毕后填充
# ============================================================

_DISPATCH.update({
    ".pdf":  _parse_pdf,
    ".docx": _parse_docx,
    ".doc":  _parse_doc,
    ".txt":  _parse_txt,
    ".png":  _parse_image,
    ".jpg":  _parse_image,
    ".jpeg": _parse_image,
    ".bmp":  _parse_image,
    ".tiff": _parse_image,
})

# ============================================================
# 兼容包装 — 供现有 job_manager.py 调用
# ============================================================

def extract_text_from_local_file(file_location: str, fallback_text: str = "") -> str:
    """
    兼容旧接口，内部委托给 parse_resume_file()。
    保留此函数确保 job_manager.py 的 4 处调用无需修改。
    """
    result = parse_resume_file(file_location, fallback_text="")
    return result if result else fallback_text


# ============================================================
# 文本分块 — 保持不变
# ============================================================

def split_text_pure_python(text: str, chunk_size: int = 250, chunk_overlap: int = 30) -> list[str]:
    """
    纯 Python 实现的文本切片器，完美替代 LangChain 避免 Windows 底层死锁
    """
    if not text:
        return []

    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = min(start + chunk_size, text_length)

        if end == text_length:
            chunks.append(text[start:end].strip())
            break

        search_start = max(start, end - chunk_overlap * 2)
        search_window = text[search_start:end]

        break_point = end
        for sep in ["\n\n", "\n", "。", "！", "？", "；", "，", "、", " "]:
            pos = search_window.rfind(sep)
            if pos != -1:
                break_point = search_start + pos + len(sep)
                break

        chunks.append(text[start:break_point].strip())

        next_start = break_point - chunk_overlap
        start = max(next_start, start + 1)

    return [c for c in chunks if c]
